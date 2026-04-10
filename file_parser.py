# -*- coding: utf-8 -*-
"""
Universal file parser for tender specifications.
Supports: html, htm, xlsx, docx, rtf, xml
Strategy: extract raw content → try direct column detection →
          fallback to GigaChat normalization.
"""

import os
import re
import io
from typing import List, Dict, Optional

from bs4 import BeautifulSoup


# ─────────────────────── DIRECT PARSERS ───────────────────────

def _clean_text(s: str) -> str:
    return re.sub(r"\s+", " ", str(s or "")).strip()


def _try_parse_number(s: str) -> float:
    s = re.sub(r"\s", "", str(s or "")).replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0


def _rows_to_items(rows: List[List[str]]) -> List[Dict]:
    """
    Try to auto-detect columns from a list of string rows.
    Returns items if a reliable structure is found, else [].
    """
    if not rows:
        return []

    # Find rows where first column is a pure integer (position number)
    data_rows = []
    for row in rows:
        if not row:
            continue
        first = _clean_text(row[0])
        if re.match(r"^\d+$", first):
            data_rows.append(row)

    if len(data_rows) < 2:
        return []

    items = []
    for row in data_rows:
        cells = [_clean_text(c) for c in row]
        # Pad to at least 7 columns
        while len(cells) < 7:
            cells.append("")

        num = int(cells[0])
        name = cells[1]
        if not name:
            continue

        # Try to find qty, unit, price from various column layouts
        qty = 0.0
        unit = "шт."
        max_price = 0.0
        desc = ""

        if len(cells) >= 5:
            qty = _try_parse_number(cells[2])
            unit = cells[3] or "шт."
            max_price = _try_parse_number(cells[4])
        if len(cells) >= 8:
            desc = cells[6]
        elif len(cells) >= 6:
            desc = cells[5]

        items.append({
            "num": num,
            "name": name,
            "qty": qty,
            "unit": unit,
            "max_price": max_price,
            "description": desc,
        })

    return items


# ─── HTML / HTM ───

def parse_html(content: bytes) -> tuple:
    """Returns (items_or_[], raw_text_for_gigachat)"""
    # Try windows-1251 first, then utf-8
    for enc in ("windows-1251", "utf-8", "cp1251"):
        try:
            text = content.decode(enc, errors="replace")
            break
        except Exception:
            text = content.decode("utf-8", errors="replace")

    soup = BeautifulSoup(text, "html.parser")
    rows_data = []
    for row in soup.find_all("tr"):
        cells = row.find_all("td")
        if cells:
            rows_data.append([c.get_text(separator=" ", strip=True) for c in cells])

    items = _rows_to_items(rows_data)

    # Build raw text for GigaChat fallback
    raw_lines = []
    for row in rows_data[:200]:
        raw_lines.append(" | ".join(row))
    raw_text = "\n".join(raw_lines)

    return items, raw_text


# ─── XLSX ───

def parse_xlsx(content: bytes) -> tuple:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    ws = wb.active

    rows_data = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c) if c is not None else "" for c in row]
        if any(c.strip() for c in cells):
            rows_data.append(cells)

    items = _rows_to_items(rows_data)

    raw_lines = []
    for row in rows_data[:200]:
        raw_lines.append(" | ".join(row))
    raw_text = "\n".join(raw_lines)

    return items, raw_text


# ─── DOCX ───

def parse_docx(content: bytes) -> tuple:
    from docx import Document
    doc = Document(io.BytesIO(content))

    rows_data = []
    raw_parts = []

    # Extract from tables first
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            if any(cells):
                rows_data.append(cells)
                raw_parts.append(" | ".join(cells))

    # Extract paragraph text as fallback
    para_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if para_text:
        raw_parts.append(para_text)

    items = _rows_to_items(rows_data)
    raw_text = "\n".join(raw_parts)
    return items, raw_text


# ─── RTF ───

def parse_rtf(content: bytes) -> tuple:
    from striprtf.striprtf import rtf_to_text
    text = rtf_to_text(content.decode("utf-8", errors="replace"))
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    rows_data = []
    for line in lines:
        # Try tab-separated
        if "\t" in line:
            rows_data.append(line.split("\t"))
        else:
            rows_data.append([line])

    items = _rows_to_items(rows_data)
    raw_text = "\n".join(lines[:200])
    return items, raw_text


# ─── XML ───

def parse_xml(content: bytes) -> tuple:
    from xml.etree import ElementTree as ET

    try:
        root = ET.fromstring(content)
    except ET.ParseError:
        # Try with encoding declaration stripped
        text = content.decode("utf-8", errors="replace")
        text = re.sub(r"<\?xml[^>]+\?>", "", text)
        root = ET.fromstring(text.encode())

    rows_data = []

    def extract_rows(node, depth=0):
        children = list(node)
        if not children:
            return
        # Check if all children have same tag (likely rows)
        tags = [c.tag for c in children]
        if len(set(tags)) <= 3 and len(children) > 1:
            for child in children:
                row_cells = []
                for sub in child.iter():
                    if sub.text and sub.text.strip():
                        row_cells.append(sub.text.strip())
                if row_cells:
                    rows_data.append(row_cells)
        else:
            for child in children:
                extract_rows(child, depth + 1)

    extract_rows(root)

    items = _rows_to_items(rows_data)

    raw_lines = []
    for row in rows_data[:200]:
        raw_lines.append(" | ".join(row))
    raw_text = "\n".join(raw_lines)

    return items, raw_text


# ─────────────────────── MAIN ENTRY POINT ───────────────────────

PARSERS = {
    ".html": parse_html,
    ".htm":  parse_html,
    ".xlsx": parse_xlsx,
    ".docx": parse_docx,
    ".rtf":  parse_rtf,
    ".xml":  parse_xml,
}


def parse_file(filepath: str, original_name: str, gigachat_token: Optional[str] = None) -> List[Dict]:
    """
    Parse a tender file and return structured items list.

    Strategy:
    1. Auto-detect format by extension
    2. Try direct column detection
    3. If items < 2 and gigachat_token is available → GigaChat normalization
    4. Filter and validate items
    """
    ext = os.path.splitext(original_name.lower())[1]
    if ext not in PARSERS:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    with open(filepath, "rb") as f:
        content = f.read()

    parser = PARSERS[ext]
    items, raw_text = parser(content)

    # If direct parsing found reasonable results — use them
    if len(items) >= 2:
        return _validate_items(items)

    # Fallback to YandexGPT normalization
    if gigachat_token and raw_text.strip():
        print(f"[Parser] Direct parse found {len(items)} items, using YandexGPT normalization...")
        from yandexgpt_client import normalize_table_to_items as yandex_normalize
        # gigachat_token here carries (api_key, folder_id) tuple
        if isinstance(gigachat_token, tuple):
            api_key, folder_id = gigachat_token
        else:
            # Legacy: treat as api_key with empty folder_id (won't work, but graceful)
            api_key, folder_id = gigachat_token, ""
        yandex_items = yandex_normalize(api_key, folder_id, raw_text)
        if yandex_items:
            return _validate_items(yandex_items)

    # Return whatever we have
    return _validate_items(items)


def _validate_items(items: List[Dict]) -> List[Dict]:
    """Filter out obviously bad rows, ensure required fields."""
    valid = []
    seen_names = set()
    for it in items:
        name = _clean_text(it.get("name", ""))
        if not name or len(name) < 2:
            continue
        # Skip header-like rows
        low = name.lower()
        if any(kw in low for kw in ("наименование", "название", "товар", "позиция", "итого", "total")):
            continue
        # Skip duplicates
        if name in seen_names:
            continue
        seen_names.add(name)

        valid.append({
            "num":         int(it.get("num", len(valid) + 1)),
            "name":        name,
            "qty":         float(it.get("qty", 0) or 0),
            "unit":        _clean_text(it.get("unit", "шт.")) or "шт.",
            "max_price":   float(it.get("max_price", 0) or 0),
            "description": _clean_text(it.get("description", "")),
        })

    # Re-number if nums are all 0
    if valid and all(it["num"] == 0 for it in valid):
        for i, it in enumerate(valid, 1):
            it["num"] = i

    return valid
